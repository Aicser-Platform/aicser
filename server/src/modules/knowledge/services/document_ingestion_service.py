"""
Document Ingestion Service for RAG Knowledge Base.

Parses uploaded documents (PDF, DOCX, MD, TXT), splits into semantic chunks,
generates embeddings via the existing EmbeddingService, and persists to the
document_chunks table.

Content coverage:
- Text: full (PDF text layer, DOCX paragraphs, MD/TXT).
- Tables: structured extraction (PDF via PyMuPDF find_tables()→Markdown; DOCX via document.tables→Markdown).
- Images / video / audio: not extracted (image-only PDF pages skipped; OCR/transcription not implemented).
See KNOWLEDGE_BASE_CONTENT_COVERAGE.md in this package for full matrix.
"""

import asyncio
import logging
import os
import re
import uuid
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

from sqlalchemy import select, update
from sqlalchemy.ext.asyncio import AsyncSession

from src.core.config import settings
from src.modules.knowledge.models import DocumentChunk, KnowledgeDocument
from src.modules.ai.utils.embedding_service import get_embedding

logger = logging.getLogger(__name__)

# ── Configuration ────────────────────────────────────────────────────────────
CHUNK_TARGET_TOKENS = 600
CHUNK_MAX_TOKENS = 800
CHUNK_OVERLAP_TOKENS = 100
EMBEDDING_BATCH_SIZE = 20
APPROX_CHARS_PER_TOKEN = 4  # rough heuristic for tiktoken-less estimation

# Max pages to process per document (0 = no limit). Prevents OOM on huge PDFs.
def _max_pages_per_document() -> int:
    return getattr(settings, "KNOWLEDGE_MAX_PAGES_PER_DOCUMENT", 0) or 0


@dataclass
class RawSection:
    """A section of parsed text with metadata."""
    text: str
    page_number: Optional[int] = None
    heading: Optional[str] = None
    source_filename: Optional[str] = None
    content_type: Optional[str] = None  # "table" | None (body text)


@dataclass
class ChunkData:
    """A text chunk ready for embedding and storage."""
    content: str
    token_count: int
    chunk_index: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)


class DocumentIngestionService:
    """
    End-to-end document ingestion: parse → chunk → embed → store.

    Usage:
        service = DocumentIngestionService(session)
        doc = await service.ingest_document(file_path, data_source_id, user_id)
    """

    def __init__(self, session: AsyncSession):
        self._session = session

    # ── Public API ───────────────────────────────────────────────────────

    async def ingest_document(
        self,
        file_path: str,
        data_source_id: str,
        user_id: str,
        filename: Optional[str] = None,
    ) -> KnowledgeDocument:
        """
        Ingest a single document end-to-end.
        Returns the KnowledgeDocument row (status will be 'ready' or 'failed').
        """
        resolved_filename = filename or os.path.basename(file_path)
        file_type = self._detect_file_type(resolved_filename)

        doc = KnowledgeDocument(
            id=uuid.uuid4(),
            data_source_id=data_source_id,
            user_id=user_id,
            filename=resolved_filename,
            file_type=file_type,
            status="processing",
        )
        self._session.add(doc)
        await self._session.commit()
        await self._session.refresh(doc)
        doc_id = doc.id

        try:
            sections, parse_meta = await self._parse_file(file_path, file_type, resolved_filename)
            if not sections:
                await self._mark_failed(doc_id, "No text content found in document")
                doc.status = "failed"
                return doc

            chunks = self._chunk_sections(sections)
            if not chunks:
                await self._mark_failed(doc_id, "Chunking produced no chunks")
                doc.status = "failed"
                return doc

            stored = await self._embed_and_store(doc_id, data_source_id, chunks)

            word_count = sum(len(s.text.split()) for s in sections)
            # For PDF: total_pages = file page count; pages_with_text = count of pages with extractable text
            total_pages = parse_meta.get("total_pages")
            pages_with_text = parse_meta.get("pages_with_text")
            if total_pages is None:
                total_pages = max((s.page_number or 0) for s in sections) if sections else 0
            if pages_with_text is None:
                pages_with_text = len(sections)

            doc_metadata: Dict[str, Any] = {
                "word_count": word_count,
                "page_count": int(total_pages),
                "pages_with_text": int(pages_with_text),
                "section_count": len(sections),
            }

            await self._session.execute(
                update(KnowledgeDocument)
                .where(KnowledgeDocument.id == doc_id)
                .values(
                    status="ready",
                    chunk_count=stored,
                    doc_metadata=doc_metadata,
                )
            )
            await self._session.commit()
            doc.status = "ready"
            doc.chunk_count = stored
            logger.info("Ingested %s: %d chunks stored", resolved_filename, stored)

        except Exception as exc:
            logger.exception("Ingestion failed for %s", resolved_filename)
            await self._mark_failed(doc_id, str(exc)[:500])
            doc.status = "failed"

        return doc

    # ── File Parsing ─────────────────────────────────────────────────────

    async def _parse_file(
        self, file_path: str, file_type: str, filename: str
    ) -> Tuple[List[RawSection], Dict[str, Any]]:
        """Route to type-specific parser. Returns (sections, parse_meta). parse_meta may include total_pages (PDF)."""
        parsers = {
            "pdf": self._parse_pdf,
            "docx": self._parse_docx,
            "md": self._parse_markdown,
            "txt": self._parse_text,
        }
        parser = parsers.get(file_type)
        if parser is None:
            raise ValueError(f"Unsupported file type: {file_type}")
        return await parser(file_path, filename)

    async def _parse_pdf(self, file_path: str, filename: str) -> Tuple[List[RawSection], Dict[str, Any]]:
        """
        Parse PDF using PyMuPDF (fitz). Processes all pages; no page limit by default.

        Only pages with extractable text (text layer) are added. Image-only or blank
        pages are skipped; for those, OCR would be needed (e.g. PyMuPDF + Tesseract).
        """
        try:
            import fitz  # pymupdf
        except ImportError:
            raise ImportError("pymupdf is required for PDF parsing: pip install pymupdf")

        loop = asyncio.get_event_loop()

        def _extract():
            sections: List[RawSection] = []
            doc = fitz.open(file_path)
            total_pages = len(doc)
            # Optional safety cap (0 = no limit)
            cap = _max_pages_per_document()
            max_pages = cap if cap > 0 else total_pages
            pages_to_process = min(total_pages, max_pages)

            for page_num in range(pages_to_process):
                page = doc[page_num]
                text = page.get_text("text")
                if text and text.strip():
                    sections.append(
                        RawSection(
                            text=text.strip(),
                            page_number=page_num + 1,
                            source_filename=filename,
                        )
                    )
                # Extract tables (PyMuPDF 1.23+): find_tables() → to_markdown() for RAG
                if hasattr(page, "find_tables"):
                    try:
                        finder = page.find_tables()
                        tbl_list = getattr(finder, "tables", None)
                        if tbl_list is None and hasattr(finder, "__iter__"):
                            tbl_list = list(finder)
                        tbl_list = tbl_list or []
                        for ti, tbl in enumerate(tbl_list):
                            if getattr(tbl, "row_count", 2) < 2 or getattr(tbl, "col_count", 2) < 2:
                                continue
                            if hasattr(tbl, "to_markdown"):
                                md = tbl.to_markdown()
                                if md and md.strip():
                                    sections.append(
                                        RawSection(
                                            text=md.strip(),
                                            page_number=page_num + 1,
                                            source_filename=filename,
                                            heading=f"Table {ti + 1}",
                                            content_type="table",
                                        )
                                    )
                    except Exception as te:  # noqa: BLE001
                        logger.debug("PDF table extraction failed on page %s: %s", page_num + 1, te)

            doc.close()

            if total_pages > 0 and len(sections) < total_pages and max_pages >= total_pages:
                logger.info(
                    "PDF %s: %d total pages, %d pages had extractable text (image-only/blank pages skipped; OCR not used)",
                    filename, total_pages, len(sections),
                )
            if cap > 0 and total_pages > max_pages:
                logger.warning(
                    "PDF %s: capped at %d pages (total %d). Set KNOWLEDGE_MAX_PAGES_PER_DOCUMENT=0 for no limit.",
                    filename, max_pages, total_pages,
                )

            parse_meta: Dict[str, Any] = {
                "total_pages": total_pages,
                "pages_with_text": len(sections),  # count of pages that had extractable text
            }
            return sections, parse_meta

        result = await loop.run_in_executor(None, _extract)
        return result

    async def _parse_docx(self, file_path: str, filename: str) -> Tuple[List[RawSection], Dict[str, Any]]:
        """Parse DOCX using python-docx. Paragraphs and tables are emitted in document order."""
        try:
            from docx import Document
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.table import Table
            from docx.text.paragraph import Paragraph
        except ImportError as e:
            raise ImportError("python-docx is required for DOCX parsing: pip install python-docx") from e

        loop = asyncio.get_event_loop()

        def _iter_blocks(document: Any) -> List[Tuple[str, Any]]:
            """Yield ('paragraph', p) or ('table', t) in document body order."""
            out: List[Tuple[str, Any]] = []
            body = document.element.body
            for child in body.iterchildren():
                if isinstance(child, CT_P):
                    out.append(("paragraph", Paragraph(child, document)))
                elif isinstance(child, CT_Tbl):
                    out.append(("table", Table(child, document)))
            return out

        def _cell_text(cell: Any) -> str:
            t = (getattr(cell, "text", None) or "").replace("|", " ").replace("\n", " ").strip()
            return t

        def _table_to_markdown(table: Any) -> str:
            rows_list: List[str] = []
            for row in table.rows:
                cells = [_cell_text(cell) for cell in row.cells]
                rows_list.append("| " + " | ".join(cells) + " |")
            if not rows_list:
                return ""
            ncols = max(1, len(rows_list[0].split("|")) - 2)
            rows_list.insert(1, "|" + " --- |" * ncols)
            return "\n".join(rows_list)

        def _extract():
            doc = Document(file_path)
            sections: List[RawSection] = []
            current_heading: Optional[str] = None
            buffer: List[str] = []
            table_idx = 0

            for block_type, block in _iter_blocks(doc):
                if block_type == "paragraph":
                    para = block
                    text = (getattr(para, "text", None) or "").strip()
                    if not text:
                        continue
                    if getattr(para, "style", None) and getattr(para.style, "name", None) and str(para.style.name).startswith("Heading"):
                        if buffer:
                            sections.append(
                                RawSection(
                                    text="\n".join(buffer),
                                    heading=current_heading,
                                    source_filename=filename,
                                )
                            )
                            buffer = []
                        current_heading = text
                    else:
                        buffer.append(text)
                else:
                    # Table: flush paragraph buffer then emit table
                    if buffer:
                        sections.append(
                            RawSection(
                                text="\n".join(buffer),
                                heading=current_heading,
                                source_filename=filename,
                            )
                        )
                        buffer = []
                    table_idx += 1
                    table_md = _table_to_markdown(block)
                    if table_md:
                        sections.append(
                            RawSection(
                                text=table_md,
                                source_filename=filename,
                                heading=f"Table {table_idx}",
                                content_type="table",
                            )
                        )

            if buffer:
                sections.append(
                    RawSection(
                        text="\n".join(buffer),
                        heading=current_heading,
                        source_filename=filename,
                    )
                )

            return sections, {}

        return await loop.run_in_executor(None, _extract)

    async def _parse_markdown(self, file_path: str, filename: str) -> Tuple[List[RawSection], Dict[str, Any]]:
        """Parse Markdown by splitting on headings."""
        loop = asyncio.get_event_loop()

        def _extract():
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()

            sections: List[RawSection] = []
            # Split on markdown headings (# ... ##... etc.)
            parts = re.split(r"(?m)^(#{1,6}\s+.+)$", content)

            current_heading: Optional[str] = None
            buffer: List[str] = []

            def _looks_like_table(text: str) -> bool:
                lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
                if len(lines) < 2:
                    return False
                has_pipe = all("|" in ln for ln in lines)
                has_sep = any(re.match(r"^[\s\|:\-]+$", ln) for ln in lines)
                return bool(has_pipe and (has_sep or len(lines) >= 2))

            for part in parts:
                part = part.strip()
                if not part:
                    continue
                if re.match(r"^#{1,6}\s+", part):
                    if buffer:
                        buf_text = "\n".join(buffer)
                        sections.append(
                            RawSection(
                                text=buf_text,
                                heading=current_heading,
                                source_filename=filename,
                                content_type="table" if _looks_like_table(buf_text) else None,
                            )
                        )
                        buffer = []
                    current_heading = re.sub(r"^#+\s*", "", part).strip()
                else:
                    buffer.append(part)

            if buffer:
                buf_text = "\n".join(buffer)
                sections.append(
                    RawSection(
                        text=buf_text,
                        heading=current_heading,
                        source_filename=filename,
                        content_type="table" if _looks_like_table(buf_text) else None,
                    )
                )
            return sections, {}

        return await loop.run_in_executor(None, _extract)

    async def _parse_text(self, file_path: str, filename: str) -> Tuple[List[RawSection], Dict[str, Any]]:
        """Parse plain text, splitting on double-newlines as paragraph boundaries."""
        loop = asyncio.get_event_loop()

        def _extract():
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()

            paragraphs = re.split(r"\n{2,}", content)
            sections: List[RawSection] = []
            for para in paragraphs:
                text = para.strip()
                if text:
                    sections.append(
                        RawSection(text=text, source_filename=filename)
                    )
            return sections, {}

        return await loop.run_in_executor(None, _extract)

    # ── Chunking ─────────────────────────────────────────────────────────

    def _chunk_sections(self, sections: List[RawSection]) -> List[ChunkData]:
        """
        Split sections into token-bounded chunks with overlap.
        Respects paragraph boundaries where possible.
        """
        chunks: List[ChunkData] = []
        idx = 0

        for section in sections:
            paragraphs = section.text.split("\n")
            buffer: List[str] = []
            buffer_tokens = 0

            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                para_tokens = self._estimate_tokens(para)

                # If single paragraph exceeds max, hard-split it
                if para_tokens > CHUNK_MAX_TOKENS:
                    if buffer:
                        chunks.append(self._make_chunk(buffer, buffer_tokens, idx, section))
                        idx += 1
                        buffer, buffer_tokens = self._overlap_tail(buffer)

                    for sub in self._hard_split(para, CHUNK_TARGET_TOKENS):
                        chunks.append(self._make_chunk([sub], self._estimate_tokens(sub), idx, section))
                        idx += 1
                    continue

                # Would adding this paragraph exceed target?
                if buffer_tokens + para_tokens > CHUNK_TARGET_TOKENS and buffer:
                    chunks.append(self._make_chunk(buffer, buffer_tokens, idx, section))
                    idx += 1
                    buffer, buffer_tokens = self._overlap_tail(buffer)

                buffer.append(para)
                buffer_tokens += para_tokens

            if buffer:
                chunks.append(self._make_chunk(buffer, buffer_tokens, idx, section))
                idx += 1

        # Re-index sequentially
        for i, chunk in enumerate(chunks):
            chunk.chunk_index = i

        return chunks

    def _make_chunk(
        self,
        paragraphs: List[str],
        token_count: int,
        idx: int,
        section: RawSection,
    ) -> ChunkData:
        meta: Dict[str, Any] = {}
        if section.page_number:
            meta["page_number"] = section.page_number
        if section.heading:
            meta["section_title"] = section.heading
        if section.content_type:
            meta["content_type"] = section.content_type
        if section.source_filename:
            meta["source_filename"] = section.source_filename

        return ChunkData(
            content="\n".join(paragraphs),
            token_count=token_count,
            chunk_index=idx,
            metadata=meta,
        )

    def _overlap_tail(self, buffer: List[str]):
        """Return the last ~CHUNK_OVERLAP_TOKENS worth of paragraphs from buffer."""
        tail: List[str] = []
        tail_tokens = 0
        for para in reversed(buffer):
            t = self._estimate_tokens(para)
            if tail_tokens + t > CHUNK_OVERLAP_TOKENS:
                break
            tail.insert(0, para)
            tail_tokens += t
        return tail, tail_tokens

    def _hard_split(self, text: str, target_tokens: int) -> List[str]:
        """Hard-split long text into ~target_tokens chunks by sentences or char boundary."""
        sentences = re.split(r"(?<=[.!?])\s+", text)
        parts: List[str] = []
        buf: List[str] = []
        buf_t = 0

        for sent in sentences:
            st = self._estimate_tokens(sent)
            if buf_t + st > target_tokens and buf:
                parts.append(" ".join(buf))
                buf = []
                buf_t = 0
            buf.append(sent)
            buf_t += st

        if buf:
            parts.append(" ".join(buf))
        return parts

    @staticmethod
    def _estimate_tokens(text: str) -> int:
        return max(1, len(text) // APPROX_CHARS_PER_TOKEN)

    # ── Embedding & Storage ──────────────────────────────────────────────

    async def _embed_and_store(
        self, doc_id: uuid.UUID, data_source_id: str, chunks: List[ChunkData]
    ) -> int:
        """Generate embeddings in batches and persist chunks to DB."""
        stored = 0

        for batch_start in range(0, len(chunks), EMBEDDING_BATCH_SIZE):
            batch = chunks[batch_start : batch_start + EMBEDDING_BATCH_SIZE]
            embeddings = await asyncio.gather(
                *(get_embedding(c.content) for c in batch),
                return_exceptions=True,
            )

            for chunk, emb_result in zip(batch, embeddings):
                embedding = None
                if isinstance(emb_result, list):
                    embedding = emb_result
                elif isinstance(emb_result, Exception):
                    logger.warning("Embedding failed for chunk %d: %s", chunk.chunk_index, emb_result)
                # get_embedding returns None on failure; chunk stored without embedding (keyword-only retrieval)

                row = DocumentChunk(
                    id=uuid.uuid4(),
                    document_id=doc_id,
                    data_source_id=data_source_id,
                    chunk_index=chunk.chunk_index,
                    content=chunk.content,
                    token_count=chunk.token_count,
                    embedding=embedding,
                    chunk_metadata=chunk.metadata or None,
                )
                self._session.add(row)
                stored += 1

            await self._session.commit()

        return stored

    # ── Helpers ──────────────────────────────────────────────────────────

    async def _mark_failed(self, doc_id: uuid.UUID, error_msg: str) -> None:
        try:
            await self._session.execute(
                update(KnowledgeDocument)
                .where(KnowledgeDocument.id == doc_id)
                .values(status="failed", error_message=error_msg[:500])
            )
            await self._session.commit()
        except Exception:
            logger.exception("Failed to mark document %s as failed", doc_id)

    @staticmethod
    def _detect_file_type(filename: str) -> str:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        mapping = {
            "pdf": "pdf",
            "docx": "docx",
            "doc": "docx",
            "md": "md",
            "markdown": "md",
            "txt": "txt",
            "text": "txt",
        }
        return mapping.get(ext, "txt")
